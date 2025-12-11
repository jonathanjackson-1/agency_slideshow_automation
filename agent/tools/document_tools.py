"""Document conversion tool for ingesting DOCX files.

This tool converts DOCX files to PDF format using docx2pdf,
enabling the agents to process Word documents.
"""

import os
import tempfile
from typing import Optional

from google.adk.tools import FunctionTool


def convert_docx_to_pdf(
    docx_path: str,
    output_path: Optional[str] = None
) -> dict:
    """Convert a DOCX file to PDF format.
    
    Args:
        docx_path: Absolute path to the input DOCX file.
        output_path: Optional path for the output PDF. If not provided,
                    the PDF will be saved alongside the DOCX with .pdf extension.
    
    Returns:
        dict with:
            - success: bool indicating if conversion succeeded
            - pdf_path: path to the generated PDF file
            - error: error message if conversion failed
    """
    try:
        from docx2pdf import convert
        
        # Validate input file exists
        if not os.path.exists(docx_path):
            return {
                "success": False,
                "pdf_path": None,
                "error": f"Input file not found: {docx_path}"
            }
        
        # Validate input is a docx file
        if not docx_path.lower().endswith('.docx'):
            return {
                "success": False,
                "pdf_path": None,
                "error": "Input file must be a .docx file"
            }
        
        # Determine output path
        if output_path is None:
            output_path = docx_path.rsplit('.', 1)[0] + '.pdf'
        
        # Convert the document
        convert(docx_path, output_path)
        
        return {
            "success": True,
            "pdf_path": output_path,
            "error": None
        }
        
    except ImportError:
        return {
            "success": False,
            "pdf_path": None,
            "error": "docx2pdf library not installed. Run: pip install docx2pdf"
        }
    except Exception as e:
        return {
            "success": False,
            "pdf_path": None,
            "error": str(e)
        }


def read_docx_text(docx_path: str) -> dict:
    """Extract text content from a DOCX file.
    
    Args:
        docx_path: Absolute path to the DOCX file.
    
    Returns:
        dict with:
            - success: bool indicating if extraction succeeded
            - text: extracted text content
            - error: error message if extraction failed
    """
    try:
        from docx import Document
        
        if not os.path.exists(docx_path):
            return {
                "success": False,
                "text": None,
                "error": f"File not found: {docx_path}"
            }
        
        doc = Document(docx_path)
        
        # Extract all paragraph text
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Extract text from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    table_text.append(" | ".join(row_text))
        
        # Combine all text
        full_text = "\n\n".join(paragraphs)
        if table_text:
            full_text += "\n\n--- Tables ---\n" + "\n".join(table_text)
        
        return {
            "success": True,
            "text": full_text,
            "error": None
        }
        
    except ImportError:
        return {
            "success": False,
            "text": None,
            "error": "python-docx library not installed. Run: pip install python-docx"
        }
    except Exception as e:
        return {
            "success": False,
            "text": None,
            "error": str(e)
        }


# Create ADK FunctionTools for use by agents
docx_to_pdf_tool = FunctionTool(func=convert_docx_to_pdf)
read_docx_tool = FunctionTool(func=read_docx_text)
